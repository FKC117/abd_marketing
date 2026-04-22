import csv
from io import BytesIO, StringIO

from .marketing_plan_schema import marketing_plan_csv_rows, marketing_plan_sections, stringify_plan_value


def _add_bullets(document, values):
    if values:
        for value in values:
            document.add_paragraph(str(value), style="List Bullet")
    else:
        document.add_paragraph("None", style="List Bullet")


def _add_key_value(document, label: str, value: str):
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value or "N/A")


def _add_marketing_plan_sections_docx(document, plan):
    for section in marketing_plan_sections(plan.output_style, plan.plan_json or {}):
        document.add_heading(section["label"], level=1)
        value = section["value"]
        if isinstance(value, dict):
            for key, item in value.items():
                _add_key_value(document, key.replace("_", " ").title(), stringify_plan_value(item) or "N/A")
        elif isinstance(value, list):
            if value and all(isinstance(item, dict) for item in value):
                for index, item in enumerate(value, start=1):
                    document.add_heading(f"{index}. {stringify_plan_value(item.get('name') or item.get('phase') or item.get('quarter') or item.get('stakeholder') or item.get('segment') or item.get('account_or_segment') or item.get('period') or item.get('horizon') or 'Item')}", level=2)
                    for item_key, item_value in item.items():
                        _add_key_value(document, item_key.replace("_", " ").title(), stringify_plan_value(item_value) or "N/A")
            else:
                _add_bullets(document, [stringify_plan_value(item) for item in value],)
        else:
            document.add_paragraph(stringify_plan_value(value) or "Not provided.")


def _add_marketing_plan_sections_pdf(story, plan, add_heading, add_body, add_bullets, subheading_style, body_style, bullet_style):
    for section in marketing_plan_sections(plan.output_style, plan.plan_json or {}):
        add_heading(story, section["label"], subheading_style)
        value = section["value"]
        if isinstance(value, dict):
            for key, item in value.items():
                add_body(story, f"{key.replace('_', ' ').title()}: {stringify_plan_value(item)}", body_style)
        elif isinstance(value, list):
            if value and all(isinstance(item, dict) for item in value):
                for index, item in enumerate(value, start=1):
                    heading = stringify_plan_value(item.get("name") or item.get("phase") or item.get("quarter") or item.get("stakeholder") or item.get("segment") or item.get("account_or_segment") or item.get("period") or item.get("horizon") or f"Item {index}")
                    add_heading(story, f"{index}. {heading}", body_style)
                    for item_key, item_value in item.items():
                        add_body(story, f"{item_key.replace('_', ' ').title()}: {stringify_plan_value(item_value)}", body_style)
            else:
                add_bullets(story, [stringify_plan_value(item) for item in value], bullet_style)
        else:
            add_body(story, stringify_plan_value(value), body_style)


def build_marketing_plan_csv(plan):
    buffer = StringIO()
    writer = csv.writer(buffer)
    for row in marketing_plan_csv_rows(plan.output_style, plan.plan_json or {}):
        writer.writerow(row)
    buffer.seek(0)
    return buffer


def build_strategy_docx(report, latest_log):
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise ValueError("python-docx is not installed. Please install requirements.txt before exporting Word files.") from exc

    def shade_cell(cell, fill: str):
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_properties.append(shading)

    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    accent = RGBColor(11, 74, 114)
    muted = RGBColor(96, 96, 96)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(report.title or "Untitled Strategy Report")
    title_run.font.color.rgb = accent

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    eyebrow_run = eyebrow.add_run("MedStratix Strategic Report")
    eyebrow_run.bold = True
    eyebrow_run.font.size = Pt(9)
    eyebrow_run.font.color.rgb = muted

    document.add_paragraph(report.executive_summary or "No executive summary available.")

    meta_table = document.add_table(rows=3, cols=2)
    meta_table.style = "Table Grid"
    meta_pairs = [
        ("Disease Focus", report.disease_focus or "All diseases"),
        ("Your Panel", f"{report.your_panel.company.name} - {report.your_panel.name}"),
        ("Competitor Panel", f"{report.competitor_panel.company.name} - {report.competitor_panel.name}"),
        ("Primary Market Account", report.market_account.name if report.market_account else "None linked"),
        ("LLM Provider", report.llm_provider or "N/A"),
        ("LLM Model", report.llm_model or "N/A"),
    ]
    for index, (label, value) in enumerate(meta_pairs):
        row = index // 2
        col = (index % 2) * 1
        cell = meta_table.cell(row, col)
        shade_cell(cell, "EAF2F8")
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.color.rgb = accent
        paragraph.add_run(value)

    document.add_paragraph("")

    your_panels = report.report_json.get("your_panels", [])
    competitor_panels = report.report_json.get("competitor_panels", [])
    if your_panels or competitor_panels:
        document.add_heading("Panel Sets", level=1)
        if your_panels:
            document.add_heading("Your Panel Set", level=2)
            for panel in your_panels:
                document.add_paragraph(
                    f"{panel.get('company', 'Unknown')} | {panel.get('name', 'Unknown')} | {panel.get('sample_type', 'N/A')}",
                    style="List Bullet",
                )
        if competitor_panels:
            document.add_heading("Competitor Panel Set", level=2)
            for panel in competitor_panels:
                document.add_paragraph(
                    f"{panel.get('company', 'Unknown')} | {panel.get('name', 'Unknown')} | {panel.get('sample_type', 'N/A')}",
                    style="List Bullet",
                )

    strategist_note = report.report_json.get("strategist_note", "")
    if strategist_note:
        document.add_heading("Strategist Note", level=1)
        document.add_paragraph(strategist_note)

    if report.report_json.get("market_accounts"):
        document.add_heading("Market Context", level=1)
        for account in report.report_json.get("market_accounts", []):
            account_paragraph = document.add_paragraph(style="List Bullet")
            account_run = account_paragraph.add_run(account.get("name", "Unknown account"))
            account_run.bold = True
            if account.get("city"):
                account_paragraph.add_run(f" | {account['city']}")
            if account.get("institution_type"):
                account_paragraph.add_run(f" | {account['institution_type']}")

    document.add_heading("SWOT", level=1)
    for title_text, items in (
        ("Strengths", report.swot_json.get("strengths", [])),
        ("Weaknesses", report.swot_json.get("weaknesses", [])),
        ("Opportunities", report.swot_json.get("opportunities", [])),
        ("Threats", report.swot_json.get("threats", [])),
    ):
        document.add_heading(title_text, level=2)
        _add_bullets(document, items)

    document.add_heading("Market Gap", level=1)
    _add_key_value(document, "Unmet Need", report.market_gap_json.get("unmet_need", "N/A"))
    _add_key_value(document, "Competitor Gap", report.market_gap_json.get("competitor_gap", "N/A"))
    _add_key_value(document, "Your Gap", report.market_gap_json.get("your_gap", "N/A"))
    _add_key_value(document, "Positioning Space", report.market_gap_json.get("positioning_space", "N/A"))

    document.add_heading("Guideline Coverage And Advantages", level=1)
    for title_text, items in (
        ("Your Advantages", report.guideline_advantages_json.get("your_advantages", [])),
        ("Competitor Advantages", report.guideline_advantages_json.get("competitor_advantages", [])),
        ("Clinical Watchouts", report.guideline_advantages_json.get("clinical_watchouts", [])),
    ):
        document.add_heading(title_text, level=2)
        _add_bullets(document, items)

    document.add_heading("Marketing Campaigns", level=1)
    campaigns = report.campaigns_json or []
    if campaigns:
        for index, campaign in enumerate(campaigns, start=1):
            document.add_heading(f"{index}. {campaign.get('name', 'Untitled Campaign')}", level=2)
            _add_key_value(document, "Audience", campaign.get("audience", "N/A"))
            _add_key_value(document, "Message", campaign.get("message", "N/A"))
            _add_key_value(document, "Channel Mix", campaign.get("channel_mix", "N/A"))
            _add_key_value(document, "Proof Point", campaign.get("proof_point", "N/A"))
            _add_key_value(document, "Call To Action", campaign.get("call_to_action", "N/A"))
    else:
        document.add_paragraph("No campaigns saved.")

    document.add_heading("Sales Pitch", level=1)
    document.add_paragraph(report.sales_pitch_text or "No sales pitch saved.")

    document.add_heading("Recommended Next Steps", level=1)
    _add_bullets(document, report.report_json.get("recommended_next_steps", []))

    if latest_log:
        document.add_section(WD_SECTION.CONTINUOUS)
        document.add_heading("LLM Audit", level=1)
        audit_table = document.add_table(rows=2, cols=3)
        audit_table.style = "Table Grid"
        audit_items = [
            ("Provider", latest_log.provider),
            ("Model", latest_log.model_name),
            ("Prompt Tokens", str(latest_log.prompt_tokens)),
            ("Response Tokens", str(latest_log.response_tokens)),
            ("Total Tokens", str(latest_log.total_tokens)),
            ("Estimated Cost USD", str(latest_log.estimated_cost_usd)),
        ]
        for index, (label, value) in enumerate(audit_items):
            cell = audit_table.cell(index // 3, index % 3)
            shade_cell(cell, "F4F7FA")
            paragraph = cell.paragraphs[0]
            label_run = paragraph.add_run(f"{label}\n")
            label_run.bold = True
            label_run.font.color.rgb = accent
            paragraph.add_run(value or "N/A")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_comparison_run_docx(run):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise ValueError("python-docx is not installed. Please install requirements.txt before exporting Word files.") from exc

    def shade_cell(cell, fill: str):
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_properties.append(shading)

    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    accent = RGBColor(11, 74, 114)
    muted = RGBColor(96, 96, 96)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(run.name or f"Comparison Run {run.pk}")
    title_run.font.color.rgb = accent

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    eyebrow_run = eyebrow.add_run("MedStratix Head-To-Head Comparison")
    eyebrow_run.bold = True
    eyebrow_run.font.size = Pt(9)
    eyebrow_run.font.color.rgb = muted

    summary = run.summary_json or {}
    document.add_paragraph(
        "This report captures the saved comparison run, including your selected panel set, "
        "the competitor panel set, and any summary metadata stored with the run."
    )

    meta_table = document.add_table(rows=2, cols=2)
    meta_table.style = "Table Grid"
    meta_pairs = [
        ("Run Name", run.name or f"Comparison Run {run.pk}"),
        ("Created By", getattr(run.created_by, "username", "System")),
        ("Disease Filter", run.disease_filter or "All diseases"),
        ("Linked Marketing Plans", str(run.marketing_plans.count())),
    ]
    for index, (label, value) in enumerate(meta_pairs):
        cell = meta_table.cell(index // 2, index % 2)
        shade_cell(cell, "EAF2F8")
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.color.rgb = accent
        paragraph.add_run(value)

    document.add_heading("Your Panels", level=1)
    for panel in run.your_panels.all():
        document.add_paragraph(
            f"{panel.company.name} | {panel.name} | {panel.get_sample_type_display()} | Price BDT: {panel.price or 'N/A'} | TAT: {panel.tat or 'N/A'}",
            style="List Bullet",
        )

    document.add_heading("Competitor Panels", level=1)
    for panel in run.competitor_panels.all():
        document.add_paragraph(
            f"{panel.company.name} | {panel.name} | {panel.get_sample_type_display()} | Price BDT: {panel.price or 'N/A'} | TAT: {panel.tat or 'N/A'}",
            style="List Bullet",
        )

    if summary:
        document.add_heading("Stored Summary", level=1)
        for key, value in summary.items():
            document.add_paragraph(f"{key}: {value}")

    if run.marketing_plans.exists():
        document.add_heading("Linked Marketing Plans", level=1)
        for plan in run.marketing_plans.all():
            document.add_paragraph(f"{plan.title} | {plan.output_style} | {plan.created_at}", style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_marketing_plan_docx(plan, latest_log):
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise ValueError("python-docx is not installed. Please install requirements.txt before exporting Word files.") from exc

    def shade_cell(cell, fill: str):
        cell_properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell_properties.append(shading)

    def add_text_block(document, title_text: str, value: str):
        document.add_heading(title_text, level=2)
        document.add_paragraph(value or "Not provided.")

    payload = plan.plan_json or {}
    sales_expectation = dict((plan.report_json or {}).get("sales_expectation", {}) or {})
    if payload.get("structure_version") == 2:
        document = Document()
        section = document.sections[0]
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

        styles = document.styles
        styles["Normal"].font.name = "Aptos"
        styles["Normal"].font.size = Pt(10.5)
        styles["Title"].font.name = "Aptos Display"
        styles["Title"].font.size = Pt(24)
        styles["Title"].font.bold = True
        styles["Heading 1"].font.name = "Aptos Display"
        styles["Heading 1"].font.size = Pt(15)
        styles["Heading 2"].font.name = "Aptos Display"
        styles["Heading 2"].font.size = Pt(12)

        accent = RGBColor(11, 74, 114)
        muted = RGBColor(96, 96, 96)

        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title.add_run(plan.title or "Untitled Marketing Plan")
        title_run.font.color.rgb = accent

        eyebrow = document.add_paragraph()
        eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
        eyebrow_run = eyebrow.add_run("MedStratix Oncology Marketing Plan")
        eyebrow_run.bold = True
        eyebrow_run.font.size = Pt(9)
        eyebrow_run.font.color.rgb = muted

        document.add_paragraph(payload.get("narrative_summary") or plan.executive_summary or "No executive summary available.")

        meta_table = document.add_table(rows=2, cols=3)
        meta_table.style = "Table Grid"
        meta_pairs = [
            ("Output Style", payload.get("plan_type_label", plan.output_style or "N/A")),
            ("Geography", plan.geography or "Not specified"),
            ("Disease Focus", plan.disease_focus or "Broad oncology"),
            ("Include Product Context", "Yes" if plan.include_product_context else "No"),
            ("LLM Provider", plan.llm_provider or "N/A"),
            ("LLM Model", plan.llm_model or "N/A"),
        ]
        for index, (label, value) in enumerate(meta_pairs):
            cell = meta_table.cell(index // 3, index % 3)
            shade_cell(cell, "EAF2F8")
            paragraph = cell.paragraphs[0]
            label_run = paragraph.add_run(f"{label}\n")
            label_run.bold = True
            label_run.font.color.rgb = accent
            paragraph.add_run(str(value or "N/A"))

        if plan.strategist_note:
            document.add_heading("Strategist Note", level=1)
            document.add_paragraph(plan.strategist_note)

        if plan.report_json.get("market_accounts"):
            document.add_heading("Market Context", level=1)
            for account in plan.report_json.get("market_accounts", []):
                bits = [account.get("name", "Unknown")]
                if account.get("city"):
                    bits.append(account["city"])
                if account.get("institution_type"):
                    bits.append(account["institution_type"])
                document.add_paragraph(" | ".join(bits), style="List Bullet")

        if any(sales_expectation.values()):
            document.add_heading("Sales Guardrails", level=1)
            _add_key_value(document, "Planning Horizon", stringify_plan_value(sales_expectation.get("planning_horizon")) or "N/A")
            _add_key_value(document, "Expected Monthly Samples", stringify_plan_value(sales_expectation.get("expected_monthly_samples")) or "N/A")
            _add_key_value(document, "Expected Quarterly Revenue (BDT)", stringify_plan_value(sales_expectation.get("expected_quarterly_revenue_bdt")) or "N/A")
            _add_key_value(document, "Expected Year-One Revenue (BDT)", stringify_plan_value(sales_expectation.get("expected_year_one_revenue_bdt")) or "N/A")
            _add_key_value(document, "Revenue Guardrail Note", stringify_plan_value(sales_expectation.get("revenue_guardrail_note")) or "N/A")

        _add_marketing_plan_sections_docx(document, plan)

        if latest_log:
            document.add_section(WD_SECTION.CONTINUOUS)
            document.add_heading("LLM Audit", level=1)
            audit_table = document.add_table(rows=2, cols=3)
            audit_table.style = "Table Grid"
            audit_items = [
                ("Provider", latest_log.provider),
                ("Model", latest_log.model_name),
                ("Prompt Tokens", str(latest_log.prompt_tokens)),
                ("Response Tokens", str(latest_log.response_tokens)),
                ("Total Tokens", str(latest_log.total_tokens)),
                ("Estimated Cost USD", str(latest_log.estimated_cost_usd)),
            ]
            for index, (label, value) in enumerate(audit_items):
                cell = audit_table.cell(index // 3, index % 3)
                shade_cell(cell, "F4F7FA")
                paragraph = cell.paragraphs[0]
                label_run = paragraph.add_run(f"{label}\n")
                label_run.bold = True
                label_run.font.color.rgb = accent
                paragraph.add_run(value or "N/A")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    executive = payload.get("executive_summary", {}) or {}
    market_research = payload.get("market_research", {}) or {}
    uvp = payload.get("unique_value_proposition", {}) or {}
    product = payload.get("product_pricing_strategy", {}) or {}
    promo = payload.get("promotional_channel_strategy", {}) or {}
    compliance = payload.get("sales_compliance_plan", {}) or {}
    forecast = payload.get("sales_targets_forecast", {}) or {}
    kpis = payload.get("follow_up_control_kpis", {}) or {}
    sales_pitch = payload.get("sales_pitch", {}) or {}

    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    accent = RGBColor(11, 74, 114)
    muted = RGBColor(96, 96, 96)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(plan.title or "Untitled Marketing Plan")
    title_run.font.color.rgb = accent

    eyebrow = document.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    eyebrow_run = eyebrow.add_run("MedStratix Oncology Marketing Plan")
    eyebrow_run.bold = True
    eyebrow_run.font.size = Pt(9)
    eyebrow_run.font.color.rgb = muted

    document.add_paragraph(executive.get("summary") or plan.executive_summary or "No executive summary available.")

    meta_table = document.add_table(rows=2, cols=3)
    meta_table.style = "Table Grid"
    meta_pairs = [
        ("Output Style", plan.output_style or "N/A"),
        ("Geography", plan.geography or "Not specified"),
        ("Disease Focus", plan.disease_focus or "Broad oncology"),
        ("Include Product Context", "Yes" if plan.include_product_context else "No"),
        ("LLM Provider", plan.llm_provider or "N/A"),
        ("LLM Model", plan.llm_model or "N/A"),
    ]
    for index, (label, value) in enumerate(meta_pairs):
        cell = meta_table.cell(index // 3, index % 3)
        shade_cell(cell, "EAF2F8")
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.color.rgb = accent
        paragraph.add_run(str(value or "N/A"))

    if plan.strategist_note:
        document.add_heading("Strategist Note", level=1)
        document.add_paragraph(plan.strategist_note)

    if plan.report_json.get("market_accounts"):
        document.add_heading("Market Context", level=1)
        for account in plan.report_json.get("market_accounts", []):
            bits = [account.get("name", "Unknown")]
            if account.get("city"):
                bits.append(account["city"])
            if account.get("institution_type"):
                bits.append(account["institution_type"])
            document.add_paragraph(" | ".join(bits), style="List Bullet")

    document.add_heading("Executive Summary", level=1)
    _add_key_value(document, "Mission", executive.get("mission", "N/A"))
    _add_key_value(document, "Value Gap", executive.get("value_gap", "N/A"))
    _add_key_value(document, "Core Goal", executive.get("core_goal", "N/A"))
    _add_key_value(document, "Summary", executive.get("summary", plan.executive_summary or "N/A"))

    document.add_heading("Market Research", level=1)
    _add_key_value(document, "Market Landscape", market_research.get("market_landscape", "N/A"))
    _add_key_value(document, "Competitor Audit", market_research.get("competitor_audit", "N/A"))
    _add_key_value(document, "Key Constraints", market_research.get("key_constraints", "N/A"))
    _add_key_value(document, "Market Distortion", market_research.get("market_distortion", "N/A"))
    _add_key_value(document, "Opportunity Map", market_research.get("opportunity_map", "N/A"))

    document.add_heading("SWOT", level=1)
    for title_text, items in (
        ("Strengths", (payload.get("swot", {}) or {}).get("strengths", [])),
        ("Weaknesses", (payload.get("swot", {}) or {}).get("weaknesses", [])),
        ("Opportunities", (payload.get("swot", {}) or {}).get("opportunities", [])),
        ("Threats", (payload.get("swot", {}) or {}).get("threats", [])),
    ):
        document.add_heading(title_text, level=2)
        _add_bullets(document, items)

    document.add_heading("Target Audience & Personas", level=1)
    personas = payload.get("target_audience_personas", []) or []
    if personas:
        for index, persona in enumerate(personas, start=1):
            document.add_heading(f"{index}. {persona.get('persona', 'Untitled Persona')}", level=2)
            _add_key_value(document, "Role", persona.get("role", "N/A"))
            _add_key_value(document, "Priority", persona.get("priority", "N/A"))
            _add_key_value(document, "Motivations", persona.get("motivations", "N/A"))
            _add_key_value(document, "Barriers", persona.get("barriers", "N/A"))
            _add_key_value(document, "Engagement Approach", persona.get("engagement_approach", "N/A"))
    else:
        document.add_paragraph("No personas saved.")

    document.add_heading("Unique Value Proposition", level=1)
    _add_key_value(document, "Headline", uvp.get("headline", "N/A"))
    _add_key_value(document, "Proof Points", uvp.get("proof_points", "N/A"))
    _add_key_value(document, "Why Now", uvp.get("why_now", "N/A"))

    document.add_heading("Product & Pricing Strategy", level=1)
    _add_key_value(document, "Portfolio Strategy", product.get("portfolio_strategy", "N/A"))
    _add_key_value(document, "Pricing Logic", product.get("pricing_logic", "N/A"))
    _add_key_value(document, "Access Strategy", product.get("access_strategy", "N/A"))
    _add_key_value(document, "Premium Justification", product.get("premium_justification", "N/A"))

    document.add_heading("Promotional & Channel Strategy", level=1)
    _add_key_value(document, "Medical Affairs", promo.get("medical_affairs", "N/A"))
    _add_key_value(document, "Academic Partnerships", promo.get("academic_partnerships", "N/A"))
    _add_key_value(document, "Digital Content", promo.get("digital_content", "N/A"))
    _add_key_value(document, "Field Activation", promo.get("field_activation", "N/A"))
    _add_key_value(document, "Channel Mix Summary", promo.get("channel_mix_summary", "N/A"))

    document.add_heading("Sales & Compliance Plan", level=1)
    _add_key_value(document, "Anti-Corruption Play", compliance.get("anti_corruption_play", "N/A"))
    _add_key_value(document, "Institutional Strategy", compliance.get("institutional_strategy", "N/A"))
    _add_key_value(document, "Logistics Management", compliance.get("logistics_management", "N/A"))
    _add_key_value(document, "Objection Handling", compliance.get("objection_handling", "N/A"))
    _add_key_value(document, "Compliance Guardrails", compliance.get("compliance_guardrails", "N/A"))

    document.add_heading("Sales Targets & Forecast", level=1)
    _add_key_value(document, "Year One Volume Targets", forecast.get("year_one_volume_targets", "N/A"))
    _add_key_value(document, "Revenue Projection Logic", forecast.get("revenue_projection_logic", "N/A"))
    _add_key_value(document, "Conversion Metrics", forecast.get("conversion_metrics", "N/A"))

    document.add_heading("Follow-up, Control & KPIs", level=1)
    _add_key_value(document, "Adoption Rate", kpis.get("adoption_rate", "N/A"))
    _add_key_value(document, "Retention Rate", kpis.get("retention_rate", "N/A"))
    _add_key_value(document, "Clinical Impact", kpis.get("clinical_impact", "N/A"))
    _add_key_value(document, "Account Growth", kpis.get("account_growth", "N/A"))
    _add_key_value(document, "Campaign Effectiveness", kpis.get("campaign_effectiveness", "N/A"))

    document.add_heading("Marketing Campaigns", level=1)
    campaigns = payload.get("campaign_plan", []) or []
    if campaigns:
        for index, campaign in enumerate(campaigns, start=1):
            document.add_heading(f"{index}. {campaign.get('name', 'Untitled Campaign')}", level=2)
            _add_key_value(document, "Audience", campaign.get("audience", "N/A"))
            _add_key_value(document, "Objective", campaign.get("objective", "N/A"))
            _add_key_value(document, "Message", campaign.get("message", "N/A"))
            _add_key_value(document, "Channel Mix", campaign.get("channel_mix", "N/A"))
            _add_key_value(document, "Timeline", campaign.get("timeline", "N/A"))
            _add_key_value(document, "Call To Action", campaign.get("call_to_action", "N/A"))
            _add_key_value(document, "KPI", campaign.get("kpi", "N/A"))
    else:
        document.add_paragraph("No campaigns saved.")

    document.add_heading("Sales Pitch", level=1)
    _add_key_value(document, "Elevator Pitch", sales_pitch.get("elevator_pitch", "N/A"))
    _add_key_value(document, "Clinician Pitch", sales_pitch.get("clinician_pitch", "N/A"))
    _add_key_value(document, "Institution Pitch", sales_pitch.get("institution_pitch", "N/A"))

    document.add_heading("Recommended Next Steps", level=1)
    _add_bullets(document, payload.get("recommended_next_steps", []))

    if latest_log:
        document.add_section(WD_SECTION.CONTINUOUS)
        document.add_heading("LLM Audit", level=1)
        audit_table = document.add_table(rows=2, cols=3)
        audit_table.style = "Table Grid"
        audit_items = [
            ("Provider", latest_log.provider),
            ("Model", latest_log.model_name),
            ("Prompt Tokens", str(latest_log.prompt_tokens)),
            ("Response Tokens", str(latest_log.response_tokens)),
            ("Total Tokens", str(latest_log.total_tokens)),
            ("Estimated Cost USD", str(latest_log.estimated_cost_usd)),
        ]
        for index, (label, value) in enumerate(audit_items):
            cell = audit_table.cell(index // 3, index % 3)
            shade_cell(cell, "F4F7FA")
            paragraph = cell.paragraphs[0]
            label_run = paragraph.add_run(f"{label}\n")
            label_run.bold = True
            label_run.font.color.rgb = accent
            paragraph.add_run(value or "N/A")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_marketing_plan_pdf(plan, latest_log):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ModuleNotFoundError as exc:
        raise ValueError("reportlab is not installed. Please install requirements.txt before exporting PDF files.") from exc

    payload = plan.plan_json or {}
    sales_expectation = dict((plan.report_json or {}).get("sales_expectation", {}) or {})
    if payload.get("structure_version") == 2:
        def esc(value):
            text = str(value or "Not provided.")
            return (
                text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )

        def add_heading(story, text, style):
            story.append(Paragraph(esc(text), style))
            story.append(Spacer(1, 0.12 * inch))

        def add_body(story, text, style):
            story.append(Paragraph(esc(text), style))
            story.append(Spacer(1, 0.12 * inch))

        def add_bullets(story, values, bullet_style):
            if values:
                for value in values:
                    story.append(Paragraph(f"â€¢ {esc(value)}", bullet_style))
            else:
                story.append(Paragraph("â€¢ None", bullet_style))
            story.append(Spacer(1, 0.12 * inch))

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("MPTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, textColor=colors.HexColor("#0b4a72"), alignment=TA_LEFT, spaceAfter=10)
        eyebrow_style = ParagraphStyle("MPEyebrow", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9, textColor=colors.HexColor("#6b7280"), spaceAfter=8)
        heading_style = ParagraphStyle("MPHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor("#0b4a72"), spaceBefore=10, spaceAfter=6)
        subheading_style = ParagraphStyle("MPSubHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, textColor=colors.HexColor("#0b4a72"), spaceBefore=8, spaceAfter=4)
        body_style = ParagraphStyle("MPBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.8, leading=13, textColor=colors.HexColor("#1f2937"))
        bullet_style = ParagraphStyle("MPBullet", parent=body_style, leftIndent=12, firstLineIndent=0, spaceAfter=2)

        story = []
        story.append(Paragraph(esc(plan.title or "Untitled Marketing Plan"), title_style))
        story.append(Paragraph("MedStratix Oncology Marketing Plan", eyebrow_style))
        story.append(Paragraph(esc(payload.get("narrative_summary") or plan.executive_summary or "No executive summary available."), body_style))
        story.append(Spacer(1, 0.16 * inch))

        meta_data = [
            ["Output Style", payload.get("plan_type_label", plan.output_style or "N/A"), "Geography", plan.geography or "Not specified"],
            ["Disease Focus", plan.disease_focus or "Broad oncology", "Model", plan.llm_model or "N/A"],
            ["Product Context", "Yes" if plan.include_product_context else "No", "Provider", plan.llm_provider or "N/A"],
        ]
        meta_table = Table(meta_data, colWidths=[1.2 * inch, 1.65 * inch, 1.2 * inch, 1.65 * inch])
        meta_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF2F8")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#9DB7CA")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C8D8E4")),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.16 * inch))

        if plan.strategist_note:
            add_heading(story, "Strategist Note", heading_style)
            add_body(story, plan.strategist_note, body_style)

        if plan.report_json.get("market_accounts"):
            add_heading(story, "Market Context", heading_style)
            for account in plan.report_json.get("market_accounts", []):
                bits = [account.get("name", "Unknown")]
                if account.get("city"):
                    bits.append(account["city"])
                if account.get("institution_type"):
                    bits.append(account["institution_type"])
                story.append(Paragraph(f"â€¢ {esc(' | '.join(bits))}", bullet_style))
            story.append(Spacer(1, 0.12 * inch))

        if any(sales_expectation.values()):
            add_heading(story, "Sales Guardrails", heading_style)
            add_body(story, f"Planning Horizon: {stringify_plan_value(sales_expectation.get('planning_horizon'))}", body_style)
            add_body(story, f"Expected Monthly Samples: {stringify_plan_value(sales_expectation.get('expected_monthly_samples'))}", body_style)
            add_body(story, f"Expected Quarterly Revenue (BDT): {stringify_plan_value(sales_expectation.get('expected_quarterly_revenue_bdt'))}", body_style)
            add_body(story, f"Expected Year-One Revenue (BDT): {stringify_plan_value(sales_expectation.get('expected_year_one_revenue_bdt'))}", body_style)
            add_body(story, f"Revenue Guardrail Note: {stringify_plan_value(sales_expectation.get('revenue_guardrail_note'))}", body_style)

        _add_marketing_plan_sections_pdf(story, plan, add_heading, add_body, add_bullets, heading_style, body_style, bullet_style)

        if latest_log:
            add_heading(story, "LLM Audit", heading_style)
            audit_data = [
                ["Provider", latest_log.provider or "N/A", "Model", latest_log.model_name or "N/A"],
                ["Prompt Tokens", str(latest_log.prompt_tokens), "Response Tokens", str(latest_log.response_tokens)],
                ["Total Tokens", str(latest_log.total_tokens), "Estimated Cost USD", str(latest_log.estimated_cost_usd)],
            ]
            audit_table = Table(audit_data, colWidths=[1.4 * inch, 1.5 * inch, 1.5 * inch, 1.6 * inch])
            audit_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F7FA")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D3DCE6")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0E7EF")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(audit_table)

        doc.build(story)
        buffer.seek(0)
        return buffer

    executive = payload.get("executive_summary", {}) or {}
    market_research = payload.get("market_research", {}) or {}
    uvp = payload.get("unique_value_proposition", {}) or {}
    product = payload.get("product_pricing_strategy", {}) or {}
    promo = payload.get("promotional_channel_strategy", {}) or {}
    compliance = payload.get("sales_compliance_plan", {}) or {}
    forecast = payload.get("sales_targets_forecast", {}) or {}
    kpis = payload.get("follow_up_control_kpis", {}) or {}
    sales_pitch = payload.get("sales_pitch", {}) or {}

    def esc(value):
        text = str(value or "Not provided.")
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    def add_heading(story, text, style):
        story.append(Paragraph(esc(text), style))
        story.append(Spacer(1, 0.12 * inch))

    def add_body(story, text, style):
        story.append(Paragraph(esc(text), style))
        story.append(Spacer(1, 0.12 * inch))

    def add_bullets(story, values, bullet_style):
        if values:
            for value in values:
                story.append(Paragraph(f"• {esc(value)}", bullet_style))
        else:
            story.append(Paragraph("• None", bullet_style))
        story.append(Spacer(1, 0.12 * inch))

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("MPTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, textColor=colors.HexColor("#0b4a72"), alignment=TA_LEFT, spaceAfter=10)
    eyebrow_style = ParagraphStyle("MPEyebrow", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=9, textColor=colors.HexColor("#6b7280"), spaceAfter=8)
    heading_style = ParagraphStyle("MPHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor("#0b4a72"), spaceBefore=10, spaceAfter=6)
    subheading_style = ParagraphStyle("MPSubHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, textColor=colors.HexColor("#0b4a72"), spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle("MPBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.8, leading=13, textColor=colors.HexColor("#1f2937"))
    bullet_style = ParagraphStyle("MPBullet", parent=body_style, leftIndent=12, firstLineIndent=0, spaceAfter=2)

    story = []
    story.append(Paragraph(esc(plan.title or "Untitled Marketing Plan"), title_style))
    story.append(Paragraph("MedStratix Oncology Marketing Plan", eyebrow_style))
    story.append(Paragraph(esc(executive.get("summary") or plan.executive_summary or "No executive summary available."), body_style))
    story.append(Spacer(1, 0.16 * inch))

    meta_data = [
        ["Output Style", plan.output_style or "N/A", "Geography", plan.geography or "Not specified"],
        ["Disease Focus", plan.disease_focus or "Broad oncology", "Model", plan.llm_model or "N/A"],
        ["Product Context", "Yes" if plan.include_product_context else "No", "Provider", plan.llm_provider or "N/A"],
    ]
    meta_table = Table(meta_data, colWidths=[1.2 * inch, 1.65 * inch, 1.2 * inch, 1.65 * inch])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF2F8")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#9DB7CA")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C8D8E4")),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.16 * inch))

    if plan.strategist_note:
        add_heading(story, "Strategist Note", heading_style)
        add_body(story, plan.strategist_note, body_style)

    if plan.report_json.get("market_accounts"):
        add_heading(story, "Market Context", heading_style)
        for account in plan.report_json.get("market_accounts", []):
            bits = [account.get("name", "Unknown")]
            if account.get("city"):
                bits.append(account["city"])
            if account.get("institution_type"):
                bits.append(account["institution_type"])
            story.append(Paragraph(f"• {esc(' | '.join(bits))}", bullet_style))
        story.append(Spacer(1, 0.12 * inch))

    add_heading(story, "Executive Summary", heading_style)
    add_body(story, f"Mission: {executive.get('mission', 'N/A')}", body_style)
    add_body(story, f"Value Gap: {executive.get('value_gap', 'N/A')}", body_style)
    add_body(story, f"Core Goal: {executive.get('core_goal', 'N/A')}", body_style)

    add_heading(story, "Market Research", heading_style)
    add_body(story, f"Market Landscape: {market_research.get('market_landscape', 'N/A')}", body_style)
    add_body(story, f"Competitor Audit: {market_research.get('competitor_audit', 'N/A')}", body_style)
    add_body(story, f"Key Constraints: {market_research.get('key_constraints', 'N/A')}", body_style)
    add_body(story, f"Market Distortion: {market_research.get('market_distortion', 'N/A')}", body_style)
    add_body(story, f"Opportunity Map: {market_research.get('opportunity_map', 'N/A')}", body_style)

    add_heading(story, "SWOT", heading_style)
    add_heading(story, "Strengths", subheading_style)
    add_bullets(story, (payload.get("swot", {}) or {}).get("strengths", []), bullet_style)
    add_heading(story, "Weaknesses", subheading_style)
    add_bullets(story, (payload.get("swot", {}) or {}).get("weaknesses", []), bullet_style)
    add_heading(story, "Opportunities", subheading_style)
    add_bullets(story, (payload.get("swot", {}) or {}).get("opportunities", []), bullet_style)
    add_heading(story, "Threats", subheading_style)
    add_bullets(story, (payload.get("swot", {}) or {}).get("threats", []), bullet_style)

    add_heading(story, "Target Audience & Personas", heading_style)
    personas = payload.get("target_audience_personas", []) or []
    if personas:
        for persona in personas:
            add_heading(story, persona.get("persona", "Untitled Persona"), subheading_style)
            add_body(story, f"Role: {persona.get('role', 'N/A')}", body_style)
            add_body(story, f"Priority: {persona.get('priority', 'N/A')}", body_style)
            add_body(story, f"Motivations: {persona.get('motivations', 'N/A')}", body_style)
            add_body(story, f"Barriers: {persona.get('barriers', 'N/A')}", body_style)
            add_body(story, f"Engagement Approach: {persona.get('engagement_approach', 'N/A')}", body_style)
    else:
        add_body(story, "No personas saved.", body_style)

    add_heading(story, "Unique Value Proposition", heading_style)
    add_body(story, f"Headline: {uvp.get('headline', 'N/A')}", body_style)
    add_body(story, f"Proof Points: {uvp.get('proof_points', 'N/A')}", body_style)
    add_body(story, f"Why Now: {uvp.get('why_now', 'N/A')}", body_style)

    add_heading(story, "Product & Pricing Strategy", heading_style)
    add_body(story, f"Portfolio Strategy: {product.get('portfolio_strategy', 'N/A')}", body_style)
    add_body(story, f"Pricing Logic: {product.get('pricing_logic', 'N/A')}", body_style)
    add_body(story, f"Access Strategy: {product.get('access_strategy', 'N/A')}", body_style)
    add_body(story, f"Premium Justification: {product.get('premium_justification', 'N/A')}", body_style)

    add_heading(story, "Promotional & Channel Strategy", heading_style)
    add_body(story, f"Medical Affairs: {promo.get('medical_affairs', 'N/A')}", body_style)
    add_body(story, f"Academic Partnerships: {promo.get('academic_partnerships', 'N/A')}", body_style)
    add_body(story, f"Digital Content: {promo.get('digital_content', 'N/A')}", body_style)
    add_body(story, f"Field Activation: {promo.get('field_activation', 'N/A')}", body_style)
    add_body(story, f"Channel Mix Summary: {promo.get('channel_mix_summary', 'N/A')}", body_style)

    add_heading(story, "Sales & Compliance Plan", heading_style)
    add_body(story, f"Anti-Corruption Play: {compliance.get('anti_corruption_play', 'N/A')}", body_style)
    add_body(story, f"Institutional Strategy: {compliance.get('institutional_strategy', 'N/A')}", body_style)
    add_body(story, f"Logistics Management: {compliance.get('logistics_management', 'N/A')}", body_style)
    add_body(story, f"Objection Handling: {compliance.get('objection_handling', 'N/A')}", body_style)
    add_body(story, f"Compliance Guardrails: {compliance.get('compliance_guardrails', 'N/A')}", body_style)

    add_heading(story, "Sales Targets & Forecast", heading_style)
    add_body(story, f"Year One Volume Targets: {forecast.get('year_one_volume_targets', 'N/A')}", body_style)
    add_body(story, f"Revenue Projection Logic: {forecast.get('revenue_projection_logic', 'N/A')}", body_style)
    add_body(story, f"Conversion Metrics: {forecast.get('conversion_metrics', 'N/A')}", body_style)

    add_heading(story, "Follow-up, Control & KPIs", heading_style)
    add_body(story, f"Adoption Rate: {kpis.get('adoption_rate', 'N/A')}", body_style)
    add_body(story, f"Retention Rate: {kpis.get('retention_rate', 'N/A')}", body_style)
    add_body(story, f"Clinical Impact: {kpis.get('clinical_impact', 'N/A')}", body_style)
    add_body(story, f"Account Growth: {kpis.get('account_growth', 'N/A')}", body_style)
    add_body(story, f"Campaign Effectiveness: {kpis.get('campaign_effectiveness', 'N/A')}", body_style)

    add_heading(story, "Marketing Campaigns", heading_style)
    campaigns = payload.get("campaign_plan", []) or []
    if campaigns:
        for index, campaign in enumerate(campaigns, start=1):
            add_heading(story, f"{index}. {campaign.get('name', 'Untitled Campaign')}", subheading_style)
            add_body(story, f"Audience: {campaign.get('audience', 'N/A')}", body_style)
            add_body(story, f"Objective: {campaign.get('objective', 'N/A')}", body_style)
            add_body(story, f"Message: {campaign.get('message', 'N/A')}", body_style)
            add_body(story, f"Channel Mix: {campaign.get('channel_mix', 'N/A')}", body_style)
            add_body(story, f"Timeline: {campaign.get('timeline', 'N/A')}", body_style)
            add_body(story, f"Call To Action: {campaign.get('call_to_action', 'N/A')}", body_style)
            add_body(story, f"KPI: {campaign.get('kpi', 'N/A')}", body_style)
    else:
        add_body(story, "No campaigns saved.", body_style)

    add_heading(story, "Sales Pitch", heading_style)
    add_body(story, f"Elevator Pitch: {sales_pitch.get('elevator_pitch', 'N/A')}", body_style)
    add_body(story, f"Clinician Pitch: {sales_pitch.get('clinician_pitch', 'N/A')}", body_style)
    add_body(story, f"Institution Pitch: {sales_pitch.get('institution_pitch', 'N/A')}", body_style)

    add_heading(story, "Recommended Next Steps", heading_style)
    add_bullets(story, payload.get("recommended_next_steps", []), bullet_style)

    if latest_log:
        add_heading(story, "LLM Audit", heading_style)
        audit_data = [
            ["Provider", latest_log.provider or "N/A", "Model", latest_log.model_name or "N/A"],
            ["Prompt Tokens", str(latest_log.prompt_tokens), "Response Tokens", str(latest_log.response_tokens)],
            ["Total Tokens", str(latest_log.total_tokens), "Estimated Cost USD", str(latest_log.estimated_cost_usd)],
        ]
        audit_table = Table(audit_data, colWidths=[1.4 * inch, 1.5 * inch, 1.5 * inch, 1.6 * inch])
        audit_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F7FA")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D3DCE6")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0E7EF")),
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(audit_table)

    doc.build(story)
    buffer.seek(0)
    return buffer


def build_final_marketing_report_docx(report):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise ValueError("python-docx is not installed. Please install requirements.txt before exporting Word files.") from exc

    payload = report.report_json or {}
    ordered_sections = payload.get("ordered_plans", []) or []

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)

    accent = RGBColor(11, 74, 114)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(report.title or "Final Marketing Report")
    title_run.font.color.rgb = accent

    document.add_paragraph(report.executive_summary or payload.get("combined_summary") or "No combined summary available.")
    _add_key_value(document, "Chronology Mode", stringify_plan_value(report.chronology_mode))
    _add_key_value(document, "Included Plans", str(len(report.ordered_plan_ids or [])))

    strategist_note = payload.get("strategist_note", "")
    if strategist_note:
        document.add_heading("Strategist Note", level=1)
        document.add_paragraph(stringify_plan_value(strategist_note))

    document.add_heading("Chronology", level=1)
    for index, item in enumerate(ordered_sections, start=1):
        document.add_heading(f"{index}. {stringify_plan_value(item.get('title'))}", level=2)
        _add_key_value(document, "Plan Type", stringify_plan_value(item.get("output_style_label")))
        _add_key_value(document, "Created", stringify_plan_value(item.get("created_at")))
        _add_key_value(document, "Summary", stringify_plan_value(item.get("summary")))
        for section in item.get("sections", []):
            document.add_heading(stringify_plan_value(section.get("label") or "Section"), level=3)
            value = section.get("value")
            if isinstance(value, dict):
                for key, item_value in value.items():
                    _add_key_value(document, key.replace("_", " ").title(), stringify_plan_value(item_value) or "N/A")
            elif isinstance(value, list):
                if value and all(isinstance(entry, dict) for entry in value):
                    for entry_index, entry in enumerate(value, start=1):
                        document.add_paragraph(f"Item {entry_index}", style="List Bullet")
                        for entry_key, entry_value in entry.items():
                            _add_key_value(document, entry_key.replace("_", " ").title(), stringify_plan_value(entry_value) or "N/A")
                else:
                    _add_bullets(document, [stringify_plan_value(entry) for entry in value])
            else:
                document.add_paragraph(stringify_plan_value(value) or "Not provided.")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_final_marketing_report_pdf(report):
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
    except ModuleNotFoundError as exc:
        raise ValueError("reportlab is not installed. Please install requirements.txt before exporting PDF files.") from exc

    payload = report.report_json or {}
    ordered_sections = payload.get("ordered_plans", []) or []

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("FinalReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, textColor=colors.HexColor("#0b4a72"), alignment=TA_LEFT)
    heading_style = ParagraphStyle("FinalReportHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14, textColor=colors.HexColor("#0b4a72"))
    subheading_style = ParagraphStyle("FinalReportSubheading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, textColor=colors.HexColor("#12384f"))
    body_style = ParagraphStyle("FinalReportBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13)
    bullet_style = ParagraphStyle("FinalReportBullet", parent=body_style, leftIndent=16)

    story = []

    def add_heading(text, style):
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 0.12 * inch))

    def add_body(text, style):
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 0.08 * inch))

    def add_bullets(items):
        if not items:
            add_body("None", body_style)
            return
        flow = ListFlowable([ListItem(Paragraph(stringify_plan_value(item), bullet_style)) for item in items], bulletType="bullet")
        story.append(flow)
        story.append(Spacer(1, 0.08 * inch))

    add_heading(report.title or "Final Marketing Report", title_style)
    add_body(report.executive_summary or payload.get("combined_summary") or "No combined summary available.", body_style)
    add_body(f"Chronology Mode: {stringify_plan_value(report.chronology_mode)}", body_style)
    add_body(f"Included Plans: {len(report.ordered_plan_ids or [])}", body_style)

    strategist_note = payload.get("strategist_note", "")
    if strategist_note:
        add_heading("Strategist Note", heading_style)
        add_body(stringify_plan_value(strategist_note), body_style)

    add_heading("Chronology", heading_style)
    for index, item in enumerate(ordered_sections, start=1):
        add_heading(f"{index}. {stringify_plan_value(item.get('title'))}", subheading_style)
        add_body(f"Plan Type: {stringify_plan_value(item.get('output_style_label'))}", body_style)
        add_body(f"Created: {stringify_plan_value(item.get('created_at'))}", body_style)
        add_body(f"Summary: {stringify_plan_value(item.get('summary'))}", body_style)
        for section in item.get("sections", []):
            add_heading(stringify_plan_value(section.get("label") or "Section"), body_style)
            value = section.get("value")
            if isinstance(value, dict):
                for key, item_value in value.items():
                    add_body(f"{key.replace('_', ' ').title()}: {stringify_plan_value(item_value)}", body_style)
            elif isinstance(value, list):
                if value and all(isinstance(entry, dict) for entry in value):
                    for entry in value:
                        add_bullets([f"{entry_key.replace('_', ' ').title()}: {stringify_plan_value(entry_value)}" for entry_key, entry_value in entry.items()])
                else:
                    add_bullets(value)
            else:
                add_body(stringify_plan_value(value), body_style)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=42, leftMargin=42, topMargin=50, bottomMargin=40)
    doc.build(story)
    buffer.seek(0)
    return buffer
